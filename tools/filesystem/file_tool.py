from __future__ import annotations

import csv
import io
import json
import mimetypes
import os
import shutil
import stat
import tempfile
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import chardet
from pydantic import BaseModel, Field, field_validator

from core.logging import get_logger
from tools.base import BaseTool, PermissionManifest, ToolInput, ToolResult

log = get_logger(__name__)


DEFAULT_INPUT_DIR  = "/tmp/agent/input"
DEFAULT_OUTPUT_DIR = "/tmp/agent/output"

MAX_READ_BYTES    = 2 * 1024 * 1024    # 2MB  — plain text reads
MAX_PARSE_BYTES   = 20 * 1024 * 1024   # 20MB — binary parser input
MAX_LIST_ENTRIES  = 500                # max items returned from list
MAX_TEXT_CHARS    = 50_000             # max chars in text output
MAX_ROWS_XLSX     = 2_000              # rows per sheet from xlsx
MAX_PAGES_PDF     = 50                 # pages extracted from pdf

class FileInput(BaseModel):

    action: str = Field(
        ...,
        description=(
            "read | write | append | list | move | delete | info | parse"
        ),
    )
    path: str = Field(
        ...,
        description=(
            "File or directory path. Resolved to absolute before use. "
            "Must be inside the tool's allowed directories."
        ),
    )
    content: str = Field(
        default="",
        description="Text content for write/append actions.",
    )
    dest: str = Field(
        default="",
        description="Destination path for move action.",
    )
    encoding: str = Field(
        default="utf-8",
        description="Text encoding for write/append (default utf-8).",
    )
    create_parents: bool = Field(
        default=True,
        description="Auto-create parent dirs on write (default True).",
    )
    include_hidden: bool = Field(
        default=False,
        description="Include hidden files/dirs in list (default False).",
    )
    ocr: bool = Field(
        default=True,
        description="Run OCR on images during parse (default True).",
    )
    max_rows: int = Field(
        default=MAX_ROWS_XLSX,
        ge=1,
        le=50_000,
        description="Max rows per sheet for xlsx/csv parse.",
    )
    max_pages: int = Field(
        default=MAX_PAGES_PDF,
        ge=1,
        le=500,
        description="Max pages for pdf parse.",
    )
    truncate_chars: int = Field(
        default=MAX_TEXT_CHARS,
        ge=100,
        le=200_000,
        description="Max chars in text output.",
    )

    @field_validator("action")
    @classmethod
    def validate_action(cls, v: str) -> str:
        allowed = {"read", "write", "append", "list", "move", "delete", "info", "parse"}
        v = v.lower().strip()
        if v not in allowed:
            raise ValueError(
                f"action must be one of {sorted(allowed)}, got: {v!r}"
            )
        return v

class FileTool(BaseTool):

    name         = "filesystem"
    description  = (
        "Read, write, append, list, move, delete, stat, or parse files. "
        "action='read': read a text file, auto-detects encoding. "
        "action='write': write content to a file (creates parent dirs). "
        "action='append': append content to an existing file. "
        "action='list': list directory with size, mime, and timestamps. "
        "action='move': rename or move a file within allowed paths. "
        "action='delete': delete a single file (not directories). "
        "action='info': get size, mime type, permissions, and timestamps. "
        "action='parse': extract structured content from PDF, DOCX, XLSX, "
        "CSV, TSV, plain text, or images (OCR). "
        "All paths must be within the declared allowed directories."
    )
    input_schema = FileInput

    def __init__(
        self,
        read_dirs:   list[str] | None = None,
        write_dirs:  list[str] | None = None,
        max_read_bytes:  int = MAX_READ_BYTES,
        max_parse_bytes: int = MAX_PARSE_BYTES,
    ) -> None:
        _read  = read_dirs  or [DEFAULT_INPUT_DIR]
        _write = write_dirs or [DEFAULT_OUTPUT_DIR]

        self.manifest = PermissionManifest(
            filesystem_read=_read,
            filesystem_write=_write,
        )

        self._read_dirs  = [Path(d).resolve() for d in _read]
        self._write_dirs = [Path(d).resolve() for d in _write]
        self._all_dirs   = self._read_dirs + [
            d for d in self._write_dirs if d not in self._read_dirs
        ]
        self._max_read_bytes  = max_read_bytes
        self._max_parse_bytes = max_parse_bytes

    def _resolve_and_check(self, path_str: str, write: bool = False) -> Path:
        resolved = Path(path_str).resolve()
        allowed  = self._write_dirs if write else self._all_dirs

        for allowed_dir in allowed:
            if resolved == allowed_dir or str(resolved).startswith(
                str(allowed_dir) + os.sep
            ):
                return resolved

        mode = "write" if write else "read"
        allowed_strs = [str(d) for d in allowed]
        raise PermissionError(
            f"Path {str(resolved)!r} is outside allowed {mode} directories. "
            f"Allowed: {allowed_strs}. "
            "Ensure the path starts with one of the listed directories."
        )

    async def execute(self, input: ToolInput) -> ToolResult:
        try:
            params = self.input_schema(**input.parameters)
        except Exception as exc:
            return ToolResult(
                success=False,
                error=f"Invalid arguments: {exc}",
            )

        try:
            if params.action == "read":
                return self._read(params)
            if params.action == "write":
                return self._write(params)
            if params.action == "append":
                return self._append(params)
            if params.action == "list":
                return self._list(params)
            if params.action == "move":
                return self._move(params)
            if params.action == "delete":
                return self._delete(params)
            if params.action == "info":
                return self._info(params)
            if params.action == "parse":
                return self._parse(params)
        except PermissionError as exc:
            return ToolResult(success=False, error=str(exc))
        except FileNotFoundError as exc:
            return ToolResult(
                success=False,
                error=f"File not found: {exc.filename!r}",
            )
        except IsADirectoryError:
            return ToolResult(
                success=False,
                error=f"{params.path!r} is a directory. Use action='list' for directories.",
            )
        except OSError as exc:
            return ToolResult(success=False, error=f"OS error: {exc}")

        return ToolResult(success=False, error=f"Unknown action: {params.action!r}")

    def _read(self, params: FileInput) -> ToolResult:
        path = self._resolve_and_check(params.path, write=False)

        file_size = path.stat().st_size
        if file_size > self._max_read_bytes:
            return ToolResult(
                success=False,
                error=(
                    f"File is {file_size:,} bytes, "
                    f"exceeding the {self._max_read_bytes:,} byte read limit. "
                    "Use action='parse' for large files, which handles "
                    "truncation per page/row."
                ),
            )

        raw_bytes = path.read_bytes()
        encoding, confidence = _detect_encoding(raw_bytes)

        try:
            text = raw_bytes.decode(encoding, errors="replace")
        except (LookupError, UnicodeDecodeError):
            text = raw_bytes.decode("utf-8", errors="replace")
            encoding = "utf-8"
            confidence = 0.5

        truncated = len(text) > params.truncate_chars
        content   = text[: params.truncate_chars]

        return ToolResult(
            success=True,
            output={
                "path":       str(path),
                "content":    content,
                "size_bytes": file_size,
                "encoding":   encoding,
                "confidence": round(confidence, 2),
                "lines":      content.count("\n") + 1,
                "truncated":  truncated,
                "total_chars": len(text),
            },
        )

    def _write(self, params: FileInput) -> ToolResult:
        path = self._resolve_and_check(params.path, write=True)

        if params.create_parents:
            path.parent.mkdir(parents=True, exist_ok=True)

        content_bytes = params.content.encode(params.encoding, errors="replace")

        tmp_path = path.with_suffix(path.suffix + ".tmp")
        try:
            tmp_path.write_bytes(content_bytes)
            tmp_path.replace(path)
        except Exception:
            tmp_path.unlink(missing_ok=True)
            raise

        return ToolResult(
            success=True,
            output={
                "path":        str(path),
                "bytes_written": len(content_bytes),
                "lines_written": params.content.count("\n") + 1,
                "encoding":    params.encoding,
            },
        )

    def _append(self, params: FileInput) -> ToolResult:
        path = self._resolve_and_check(params.path, write=True)

        if params.create_parents:
            path.parent.mkdir(parents=True, exist_ok=True)

        content_bytes = params.content.encode(params.encoding, errors="replace")

        with path.open("ab") as f:
            f.write(content_bytes)

        return ToolResult(
            success=True,
            output={
                "path":            str(path),
                "bytes_appended":  len(content_bytes),
                "total_size_bytes": path.stat().st_size,
            },
        )

    def _list(self, params: FileInput) -> ToolResult:
        path = self._resolve_and_check(params.path, write=False)

        if not path.is_dir():
            return ToolResult(
                success=False,
                error=f"{str(path)!r} is not a directory. Use action='read' for files.",
            )

        entries_raw = list(path.iterdir())
        if not params.include_hidden:
            entries_raw = [e for e in entries_raw if not e.name.startswith(".")]

        entries_raw.sort(key=lambda e: (not e.is_dir(), e.name.lower()))

        entries = []
        for entry in entries_raw[: MAX_LIST_ENTRIES]:
            try:
                st   = entry.stat()
                size = st.st_size
                mtime = datetime.fromtimestamp(st.st_mtime, tz=timezone.utc).isoformat()
            except OSError:
                size  = 0
                mtime = ""

            mime, _ = mimetypes.guess_type(entry.name)

            entries.append({
                "name":        entry.name,
                "path":        str(entry),
                "is_dir":      entry.is_dir(),
                "size_bytes":  size,
                "mime_type":   mime or ("directory" if entry.is_dir() else "application/octet-stream"),
                "mtime_iso":   mtime,
                "is_readable": os.access(entry, os.R_OK),
                "is_writable": os.access(entry, os.W_OK),
                "is_hidden":   entry.name.startswith("."),
            })

        total = len(entries_raw)
        return ToolResult(
            success=True,
            output={
                "path":        str(path),
                "entries":     entries,
                "count":       len(entries),
                "total_count": total,
                "truncated":   total > MAX_LIST_ENTRIES,
            },
        )

    def _move(self, params: FileInput) -> ToolResult:
        if not params.dest:
            return ToolResult(
                success=False,
                error="action='move' requires a 'dest' parameter (destination path).",
            )

        src  = self._resolve_and_check(params.path, write=False)
        dest = self._resolve_and_check(params.dest, write=True)

        if not src.exists():
            return ToolResult(
                success=False,
                error=f"Source file not found: {str(src)!r}",
            )
        if dest.exists():
            return ToolResult(
                success=False,
                error=(
                    f"Destination already exists: {str(dest)!r}. "
                    "Delete it first or choose a different destination."
                ),
            )

        dest.parent.mkdir(parents=True, exist_ok=True)
        shutil.move(str(src), str(dest))

        return ToolResult(
            success=True,
            output={
                "source":      str(src),
                "destination": str(dest),
                "size_bytes":  dest.stat().st_size,
            },
        )

    def _delete(self, params: FileInput) -> ToolResult:
        path = self._resolve_and_check(params.path, write=True)

        if not path.exists():
            return ToolResult(
                success=False,
                error=f"File not found: {str(path)!r}",
            )
        if path.is_dir():
            return ToolResult(
                success=False,
                error=(
                    f"{str(path)!r} is a directory. "
                    "This tool does not delete directories for safety. "
                    "List the directory and delete files individually."
                ),
            )

        size = path.stat().st_size
        path.unlink()

        return ToolResult(
            success=True,
            output={
                "path":          str(path),
                "bytes_freed":   size,
                "deleted":       True,
            },
        )

    def _info(self, params: FileInput) -> ToolResult:
        path = self._resolve_and_check(params.path, write=False)

        if not path.exists():
            return ToolResult(
                success=False,
                error=f"Path not found: {str(path)!r}",
            )

        st = path.stat()
        mime, encoding_hint = mimetypes.guess_type(path.name)
        perm_octal = oct(stat.S_IMODE(st.st_mode))

        mtime = datetime.fromtimestamp(st.st_mtime, tz=timezone.utc).isoformat()
        ctime = datetime.fromtimestamp(st.st_ctime, tz=timezone.utc).isoformat()

        return ToolResult(
            success=True,
            output={
                "path":          str(path),
                "name":          path.name,
                "suffix":        path.suffix.lower(),
                "is_dir":        path.is_dir(),
                "is_symlink":    path.is_symlink(),
                "size_bytes":    st.st_size,
                "mime_type":     mime or "application/octet-stream",
                "encoding_hint": encoding_hint,
                "permissions":   perm_octal,
                "mtime_iso":     mtime,
                "ctime_iso":     ctime,
                "is_readable":   os.access(path, os.R_OK),
                "is_writable":   os.access(path, os.W_OK),
            },
        )

    def _parse(self, params: FileInput) -> ToolResult:
        path = self._resolve_and_check(params.path, write=False)

        if not path.exists():
            return ToolResult(
                success=False,
                error=f"File not found: {str(path)!r}",
            )
        if path.is_dir():
            return ToolResult(
                success=False,
                error="Cannot parse a directory. Use action='list' then parse individual files.",
            )

        file_size = path.stat().st_size
        if file_size > self._max_parse_bytes:
            return ToolResult(
                success=False,
                error=(
                    f"File is {file_size:,} bytes, exceeding the "
                    f"{self._max_parse_bytes:,} byte parse limit."
                ),
            )

        suffix = path.suffix.lower()

        if suffix == ".pdf":
            return self._parse_pdf(path, params)
        if suffix in (".docx", ".doc"):
            return self._parse_docx(path, params)
        if suffix in (".xlsx", ".xls", ".xlsm"):
            return self._parse_xlsx(path, params)
        if suffix in (".csv", ".tsv"):
            return self._parse_csv(path, params)
        if suffix in (".png", ".jpg", ".jpeg", ".gif", ".bmp",
                      ".tiff", ".tif", ".webp", ".ico"):
            return self._parse_image(path, params)
        if suffix in (".txt", ".md", ".rst", ".log", ".py", ".js",
                      ".ts", ".json", ".xml", ".html", ".htm",
                      ".yaml", ".yml", ".toml", ".ini", ".cfg",
                      ".sh", ".bash", ".sql", ".r", ".rb", ".go",
                      ".java", ".c", ".cpp", ".h", ".cs", ".php",
                      ".swift", ".kt", ".rs"):
            return self._parse_text(path, params)

        log.debug("file_tool_parse_unknown_ext", ext=suffix, path=str(path))
        return self._parse_text(path, params)

    def _parse_pdf(self, path: Path, params: FileInput) -> ToolResult:
        try:
            from pypdf import PdfReader
        except ImportError:
            return ToolResult(
                success=False,
                error="pypdf is not installed. Run: poetry add pypdf",
            )

        try:
            reader = PdfReader(str(path))
        except Exception as exc:
            return ToolResult(
                success=False,
                error=f"Failed to open PDF {path.name!r}: {exc}",
            )

        total_pages = len(reader.pages)
        extract_pages = min(total_pages, params.max_pages)

        # PDF metadata
        meta: dict[str, str] = {}
        if reader.metadata:
            for key in ("/Title", "/Author", "/Subject", "/Creator", "/Producer"):
                val = reader.metadata.get(key, "")
                if val:
                    meta[key.lstrip("/")] = str(val)[:200]

        pages_output: list[dict[str, Any]] = []
        has_ocr_pages = False
        total_chars   = 0

        for page_num in range(extract_pages):
            page = reader.pages[page_num]
            try:
                text = page.extract_text(extraction_mode="plain") or ""
            except Exception:
                text = ""

            text = text.strip()
            ocr_used = False

            if len(text) < 30 and params.ocr:
                ocr_text = _ocr_pdf_page(page)
                if ocr_text:
                    text     = ocr_text
                    ocr_used = True
                    has_ocr_pages = True

            total_chars += len(text)
            pages_output.append({
                "page":       page_num + 1,
                "text":       text[: params.truncate_chars // max(1, extract_pages)],
                "char_count": len(text),
                "ocr_used":   ocr_used,
            })

        return ToolResult(
            success=True,
            output={
                "path":            str(path),
                "format":          "pdf",
                "total_pages":     total_pages,
                "extracted_pages": extract_pages,
                "truncated":       total_pages > extract_pages,
                "has_ocr_pages":   has_ocr_pages,
                "total_chars":     total_chars,
                "metadata":        meta,
                "pages":           pages_output,
            },
        )

    def _parse_docx(self, path: Path, params: FileInput) -> ToolResult:
        try:
            import docx
        except ImportError:
            return ToolResult(
                success=False,
                error="python-docx is not installed. Run: poetry add python-docx",
            )

        try:
            doc = docx.Document(str(path))
        except Exception as exc:
            return ToolResult(
                success=False,
                error=f"Failed to open DOCX {path.name!r}: {exc}",
            )

        paragraphs: list[dict[str, str]] = []
        full_text_parts: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name if para.style else "Normal"
            paragraphs.append({"style": style_name, "text": text})
            full_text_parts.append(text)

        tables_output: list[dict[str, Any]] = []
        for table_idx, table in enumerate(doc.tables):
            rows: list[list[str]] = []
            for row in table.rows:
                cell_texts = [cell.text.strip() for cell in row.cells]
                rows.append(cell_texts)

            headers: list[str] = []
            data_rows: list[dict[str, str]] = []
            if rows:
                headers = rows[0]
                for row in rows[1:]:
                    if len(row) == len(headers):
                        data_rows.append(dict(zip(headers, row)))
                    else:
                        data_rows.append({str(i): v for i, v in enumerate(row)})

            tables_output.append({
                "index":    table_idx,
                "headers":  headers,
                "rows":     data_rows,
                "row_count": len(data_rows),
            })

        props: dict[str, str] = {}
        cp = doc.core_properties
        for attr in ("author", "title", "subject", "description",
                     "keywords", "created", "modified"):
            val = getattr(cp, attr, None)
            if val:
                props[attr] = str(val)[:200]

        full_text = "\n".join(full_text_parts)
        truncated = len(full_text) > params.truncate_chars

        return ToolResult(
            success=True,
            output={
                "path":          str(path),
                "format":        "docx",
                "paragraph_count": len(paragraphs),
                "table_count":   len(tables_output),
                "char_count":    len(full_text),
                "truncated":     truncated,
                "properties":    props,
                "full_text":     full_text[: params.truncate_chars],
                "paragraphs":    paragraphs[:500],
                "tables":        tables_output,
            },
        )

    def _parse_xlsx(self, path: Path, params: FileInput) -> ToolResult:
        try:
            import openpyxl
        except ImportError:
            return ToolResult(
                success=False,
                error="openpyxl is not installed. Run: poetry add openpyxl",
            )

        try:
            wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
        except Exception as exc:
            return ToolResult(
                success=False,
                error=f"Failed to open XLSX {path.name!r}: {exc}",
            )

        sheets_output: list[dict[str, Any]] = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            all_rows = list(ws.iter_rows(values_only=True))

            if not all_rows:
                sheets_output.append({
                    "name":      sheet_name,
                    "headers":   [],
                    "rows":      [],
                    "row_count": 0,
                })
                continue

            raw_headers = all_rows[0]
            headers: list[str] = [
                str(h).strip() if h is not None else f"col_{i}"
                for i, h in enumerate(raw_headers)
            ]

            seen: dict[str, int] = {}
            unique_headers: list[str] = []
            for h in headers:
                if h in seen:
                    seen[h] += 1
                    unique_headers.append(f"{h}_{seen[h]}")
                else:
                    seen[h] = 0
                    unique_headers.append(h)

            data_rows: list[dict[str, Any]] = []
            for row in all_rows[1 : params.max_rows + 1]:
                row_dict: dict[str, Any] = {}
                for i, (header, cell_val) in enumerate(zip(unique_headers, row)):
                    if cell_val is None:
                        row_dict[header] = None
                    elif isinstance(cell_val, (int, float, bool)):
                        row_dict[header] = cell_val
                    else:
                        row_dict[header] = str(cell_val).strip()
                data_rows.append(row_dict)

            total_data_rows = len(all_rows) - 1
            sheets_output.append({
                "name":          sheet_name,
                "headers":       unique_headers,
                "rows":          data_rows,
                "row_count":     len(data_rows),
                "total_rows":    total_data_rows,
                "truncated":     total_data_rows > params.max_rows,
                "column_count":  len(unique_headers),
            })

        wb.close()

        return ToolResult(
            success=True,
            output={
                "path":        str(path),
                "format":      "xlsx",
                "sheet_count": len(sheets_output),
                "sheet_names": wb.sheetnames,
                "sheets":      sheets_output,
            },
        )

    def _parse_csv(self, path: Path, params: FileInput) -> ToolResult:
        raw_bytes = path.read_bytes()
        if len(raw_bytes) > self._max_parse_bytes:
            return ToolResult(
                success=False,
                error=f"CSV too large: {len(raw_bytes):,} bytes.",
            )

        encoding, _ = _detect_encoding(raw_bytes)
        try:
            text = raw_bytes.decode(encoding, errors="replace")
        except (LookupError, UnicodeDecodeError):
            text = raw_bytes.decode("utf-8", errors="replace")
            encoding = "utf-8"

        sample  = text[:4096]
        sniffer = csv.Sniffer()
        try:
            dialect = sniffer.sniff(sample)
            has_header = sniffer.has_header(sample)
        except csv.Error:
            dialect    = csv.excel       
            has_header = True

        reader = csv.reader(io.StringIO(text), dialect=dialect)
        all_rows_raw = list(reader)

        if not all_rows_raw:
            return ToolResult(
                success=True,
                output={
                    "path":      str(path),
                    "format":    "csv",
                    "headers":   [],
                    "rows":      [],
                    "row_count": 0,
                    "delimiter": repr(dialect.delimiter),
                    "encoding":  encoding,
                },
            )

        if has_header:
            headers  = [h.strip() for h in all_rows_raw[0]]
            data_raw = all_rows_raw[1:]
        else:
            num_cols = max((len(r) for r in all_rows_raw), default=0)
            headers  = [f"col_{i}" for i in range(num_cols)]
            data_raw = all_rows_raw

        seen_h: dict[str, int] = {}
        final_headers: list[str] = []
        for h in headers:
            if h in seen_h:
                seen_h[h] += 1
                final_headers.append(f"{h}_{seen_h[h]}")
            else:
                seen_h[h] = 0
                final_headers.append(h)

        rows: list[dict[str, str]] = []
        for row in data_raw[: params.max_rows]:
            row_dict: dict[str, str] = {}
            for i, h in enumerate(final_headers):
                row_dict[h] = row[i].strip() if i < len(row) else ""
            rows.append(row_dict)

        total_data = len(data_raw)

        return ToolResult(
            success=True,
            output={
                "path":          str(path),
                "format":        "csv",
                "headers":       final_headers,
                "rows":          rows,
                "row_count":     len(rows),
                "total_rows":    total_data,
                "truncated":     total_data > params.max_rows,
                "column_count":  len(final_headers),
                "delimiter":     repr(dialect.delimiter),
                "encoding":      encoding,
                "has_header":    has_header,
            },
        )


    def _parse_image(self, path: Path, params: FileInput) -> ToolResult:
        try:
            from PIL import Image as PILImage, ExifTags
        except ImportError:
            return ToolResult(
                success=False,
                error="Pillow is not installed. Run: poetry add Pillow",
            )

        try:
            img = PILImage.open(str(path))
            img.load()  
        except Exception as exc:
            return ToolResult(
                success=False,
                error=f"Failed to open image {path.name!r}: {exc}",
            )

        width, height = img.size
        fmt   = img.format or path.suffix.upper().lstrip(".")
        mode  = img.mode

        exif_data: dict[str, Any] = {}
        try:
            raw_exif = img.getexif()
            if raw_exif:
                for tag_id, value in raw_exif.items():
                    tag_name = ExifTags.TAGS.get(tag_id, f"Tag_{tag_id}")
                    # Skip binary blobs
                    if isinstance(value, bytes):
                        continue
                    # Convert IFDRational to float
                    try:
                        if hasattr(value, "numerator"):
                            value = float(value)
                    except Exception:
                        pass
                    exif_data[tag_name] = (
                        str(value)[:200] if not isinstance(value, (int, float, bool)) else value
                    )

            gps_info = raw_exif.get_ifd(0x8825)
            if gps_info:
                gps = _parse_gps_info(gps_info)
                if gps:
                    exif_data["GPS"] = gps
        except Exception:
            pass 

        # OCR
        ocr_text       = ""
        ocr_confidence = 0.0

        if params.ocr:
            ocr_text, ocr_confidence = _run_ocr(img)

        return ToolResult(
            success=True,
            output={
                "path":            str(path),
                "format":          fmt,
                "width":           width,
                "height":          height,
                "mode":            mode,
                "size_bytes":      path.stat().st_size,
                "exif":            exif_data,
                "ocr_text":        ocr_text[: params.truncate_chars],
                "ocr_confidence":  round(ocr_confidence, 2),
                "ocr_available":   _tesseract_available(),
            },
        )

    def _parse_text(self, path: Path, params: FileInput) -> ToolResult:
        result = self._read(params)
        if result.success and isinstance(result.output, dict):
            result.output["format"] = path.suffix.lower().lstrip(".") or "text"
            mime, _ = mimetypes.guess_type(path.name)
            result.output["mime_type"] = mime or "text/plain"
        return result

def _detect_encoding(raw: bytes) -> tuple[str, float]:
    try:
        result = chardet.detect(raw[:8192])  # sample first 8KB
        enc    = result.get("encoding") or "utf-8"
        conf   = result.get("confidence") or 0.0
        # Normalise encoding name
        enc = enc.lower().replace("-", "_")
        if conf < 0.6:
            enc = "utf-8"
        return enc, conf
    except Exception:
        return "utf-8", 0.5

def _ocr_pdf_page(page: Any) -> str:
    if not _tesseract_available():
        return ""
    try:
        from PIL import Image as PILImage
        import pytesseract

        images = getattr(page, "images", [])
        if not images:
            return ""

        img_data = images[0].data
        img      = PILImage.open(io.BytesIO(img_data))
        text     = pytesseract.image_to_string(img, timeout=20).strip()
        return text
    except Exception:
        return ""

def _run_ocr(img: Any) -> tuple[str, float]:
    if not _tesseract_available():
        return "", 0.0
    try:
        import pytesseract

        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")

        data   = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT, timeout=30)
        texts  = data.get("text", [])
        confs  = data.get("conf", [])

        valid = [
            (t, c) for t, c in zip(texts, confs)
            if t.strip() and isinstance(c, (int, float)) and c >= 0
        ]

        if not valid:
            return "", 0.0

        text        = " ".join(t for t, _ in valid)
        mean_conf   = sum(c for _, c in valid) / len(valid) / 100.0  # normalize 0-1
        return text.strip(), round(mean_conf, 3)

    except Exception:
        return "", 0.0

def _parse_gps_info(gps_ifd: dict) -> dict[str, Any] | None:
    try:
        from PIL.ExifTags import GPSTAGS

        def to_decimal(dms: Any) -> float:
            try:
                d, m, s = float(dms[0]), float(dms[1]), float(dms[2])
                return d + m / 60 + s / 3600
            except Exception:
                return 0.0

        gps: dict[str, Any] = {}
        for k, v in gps_ifd.items():
            tag_name = GPSTAGS.get(k, f"GPS_{k}")
            gps[tag_name] = v

        lat_dms = gps.get("GPSLatitude")
        lon_dms = gps.get("GPSLongitude")
        lat_ref = gps.get("GPSLatitudeRef", "N")
        lon_ref = gps.get("GPSLongitudeRef", "E")
        alt     = gps.get("GPSAltitude")

        if not (lat_dms and lon_dms):
            return None

        lat = to_decimal(lat_dms) * (-1 if lat_ref == "S" else 1)
        lon = to_decimal(lon_dms) * (-1 if lon_ref == "W" else 1)

        result: dict[str, Any] = {"lat": round(lat, 6), "lon": round(lon, 6)}
        if alt is not None:
            try:
                result["alt_m"] = round(float(alt), 1)
            except Exception:
                pass
        return result
    except Exception:
        return None

_tesseract_checked: bool = False
_tesseract_ok:      bool = False


def _tesseract_available() -> bool:
    global _tesseract_checked, _tesseract_ok
    if _tesseract_checked:
        return _tesseract_ok
    try:
        import subprocess as _sp
        result = _sp.run(
            ["tesseract", "--version"],
            capture_output=True,
            timeout=3,
        )
        _tesseract_ok = result.returncode == 0
    except Exception:
        _tesseract_ok = False
    _tesseract_checked = True
    return _tesseract_ok